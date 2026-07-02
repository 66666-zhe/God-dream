# -*- coding: utf-8 -*-
"""生成包含城市土地覆盖分类代码的 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ========== 标题 ==========
title = doc.add_heading('城市土地覆盖分类 —— 高斯朴素贝叶斯', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('数据集来源：UCI Machine Learning Repository — Urban Land Cover')
doc.add_paragraph('')

# ========== 代码内容 ==========
code_blocks = [
    (
        '（1）载入训练集与测试集，并分别考察其形状',
        """import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.naive_bayes import GaussianNB
from sklearn.metrics import confusion_matrix, accuracy_score

# 载入数据
train = pd.read_csv('land_train.csv')
test  = pd.read_csv('land_test.csv')

print('训练集形状：', train.shape)
print('测试集形状：', test.shape)"""
    ),
    (
        '（2）考察训练集中响应变量 Class 的分布，并画柱状图',
        """print('训练集 Class 分布：')
print(train['Class'].value_counts())

sns.catplot(x='Class', kind='count', data=train)
plt.title('Training Set - Class Distribution')
plt.xticks(rotation=45)
plt.tight_layout()
plt.show()"""
    ),
    (
        '（3）进行高斯朴素贝叶斯估计',
        """# 分离特征与标签
X_train = train.drop('Class', axis=1)
y_train = train['Class']

X_test  = test.drop('Class', axis=1)
y_test  = test['Class']

# 建立并拟合模型
gnb = GaussianNB()
gnb.fit(X_train, y_train)

print('高斯朴素贝叶斯模型训练完成。')"""
    ),
    (
        '（4）在测试集中预测分类概率，展示前 5 个预测概率',
        """y_prob = gnb.predict_proba(X_test)

print('各类别：', gnb.classes_)
print('\\n前 5 个样本的预测概率：')
print(pd.DataFrame(y_prob[:5], columns=gnb.classes_))"""
    ),
    (
        '（5）在测试集中预测分类结果，展示前 5 个预测结果',
        """y_pred = gnb.predict(X_test)

print('前 5 个预测结果：', y_pred[:5])
print('前 5 个真实标签：', y_test.values[:5])"""
    ),
    (
        '（6）展示测试集的混淆矩阵，并计算预测准确率',
        """cm = confusion_matrix(y_test, y_pred, labels=gnb.classes_)
cm_df = pd.DataFrame(cm, index=gnb.classes_, columns=gnb.classes_)

print('混淆矩阵：')
print(cm_df)

acc = accuracy_score(y_test, y_pred)
print(f'\\n预测准确率：{acc:.4f}')"""
    ),
]

for title_text, code in code_blocks:
    doc.add_heading(title_text, level=1)
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

# ========== 保存 ==========
output_path = 'D:/桌面/land/城市土地覆盖分类代码.docx'
doc.save(output_path)
print(f'Word 文档已保存至：{output_path}')
