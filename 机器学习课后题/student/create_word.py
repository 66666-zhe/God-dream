# -*- coding: utf-8 -*-
"""生成 Word 文档：惩罚回归代码"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ---------- 样式设置 ----------
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---------- 标题 ----------
title = doc.add_heading('', level=1)
run = title.add_run('9.1  惩罚回归 —— 葡萄牙高中数学成绩数据')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph('数据来源：UCI Machine Learning Repository 的 student-mat.csv\n'
                   '响应变量：G3（期末成绩）\n'
                   '运行环境：Anaconda3（已安装 scikit-learn、pandas、matplotlib、numpy）')

# ---------- 辅助函数 ----------
def add_code_block(doc, code_text, title=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)

# ============================================================
# (1)
# ============================================================
add_code_block(doc, r"""
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import (Ridge, RidgeCV, Lasso, LassoCV,
                                   ElasticNet, ElasticNetCV)
from sklearn.linear_model import lasso_path
from sklearn.model_selection import KFold, train_test_split

plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
""", "导入库")

add_code_block(doc, r"""
# (1) 载入数据，考察形状与前5个观测值
df = pd.read_table('student-mat.csv', sep=';')
print("数据形状:", df.shape)
print(df.head())
""", "(1) 载入数据")

add_code_block(doc, r"""
# (2) 去掉变量 G1 与 G2
df = df.drop(['G1', 'G2'], axis=1)
print("去掉 G1、G2 后形状:", df.shape)
""", "(2) 去掉 G1 与 G2")

add_code_block(doc, r"""
# (3) 画响应变量 G3 的直方图
plt.figure(figsize=(8, 5))
plt.hist(df['G3'], bins=20, edgecolor='black', alpha=0.7)
plt.xlabel('G3 (期末成绩)')
plt.ylabel('频数')
plt.title('响应变量 G3 的直方图')
plt.tight_layout()
plt.show()
""", "(3) G3 直方图")

add_code_block(doc, r"""
# (4) 将分类变量变为虚拟变量
df_dummies = pd.get_dummies(df, drop_first=True)
print("虚拟变量化后形状:", df_dummies.shape)
""", "(4) 虚拟变量")

add_code_block(doc, r"""
# (5) 将所有特征变量标准化
y = df_dummies['G3'].values
X = df_dummies.drop('G3', axis=1).values
feature_names = df_dummies.drop('G3', axis=1).columns.tolist()

scaler = StandardScaler()
X_std = scaler.fit_transform(X)
print("标准化后特征矩阵形状:", X_std.shape)
""", "(5) 标准化")

add_code_block(doc, r"""
# (6) 岭回归系数路径
alphas_ridge = np.logspace(-3, 6, 100)
coefs_ridge = []
for a in alphas_ridge:
    ridge = Ridge(alpha=a, fit_intercept=True)
    ridge.fit(X_std, y)
    coefs_ridge.append(ridge.coef_)
coefs_ridge = np.array(coefs_ridge)

plt.figure(figsize=(10, 6))
for i in range(coefs_ridge.shape[1]):
    plt.plot(np.log10(alphas_ridge), coefs_ridge[:, i])
plt.xlabel('log10(α)')
plt.ylabel('系数值')
plt.title('岭回归系数路径')
plt.tight_layout()
plt.show()
""", "(6) 岭回归系数路径")

add_code_block(doc, r"""
# (7) 10折交叉验证选择最优α，岭回归
kf = KFold(n_splits=10, shuffle=True, random_state=1)
ridge_cv = RidgeCV(alphas=np.logspace(-3, 6, 100), cv=kf, scoring='r2')
ridge_cv.fit(X_std, y)
print("岭回归最优 α =", ridge_cv.alpha_)

ridge_best = Ridge(alpha=ridge_cv.alpha_)
ridge_best.fit(X_std, y)
ridge_coef_df = pd.DataFrame({'特征': feature_names, '系数': ridge_best.coef_})
print(ridge_coef_df.to_string(index=False))
""", "(7) 岭回归交叉验证")

add_code_block(doc, r"""
# (8) Lasso 回归系数路径 (eps=1e-4)
alphas_lasso_path, coefs_lasso_path, _ = lasso_path(X_std, y, eps=1e-4, n_alphas=100)

plt.figure(figsize=(10, 6))
for i in range(coefs_lasso_path.shape[0]):
    plt.plot(np.log10(alphas_lasso_path), coefs_lasso_path[i, :])
plt.xlabel('log10(α)')
plt.ylabel('系数值')
plt.title('Lasso 回归系数路径 (eps=1e-4)')
plt.tight_layout()
plt.show()
""", "(8) Lasso 系数路径")

add_code_block(doc, r"""
# (9) 10折交叉验证选择最优α，Lasso 回归
lasso_cv = LassoCV(alphas=np.logspace(-3, 1, 100), cv=kf, random_state=1, max_iter=10000)
lasso_cv.fit(X_std, y)
print("Lasso 最优 α =", lasso_cv.alpha_)

lasso_best = Lasso(alpha=lasso_cv.alpha_, max_iter=10000)
lasso_best.fit(X_std, y)
lasso_coef_df = pd.DataFrame({'特征': feature_names, '系数': lasso_best.coef_})
print(lasso_coef_df.to_string(index=False))
""", "(9) Lasso 交叉验证")

add_code_block(doc, r"""
# (10) 弹性网回归：网格搜索 α 与 l1_ratio
alphas_en = np.logspace(-3, 1, 100)
l1_ratios = [0.001, 0.01, 0.1, 0.5, 1]

en_cv = ElasticNetCV(alphas=alphas_en, l1_ratio=l1_ratios,
                      cv=kf, random_state=1, max_iter=10000)
en_cv.fit(X_std, y)
print("弹性网最优 α =", en_cv.alpha_)
print("最优 l1_ratio =", en_cv.l1_ratio_)
print("样本内 R² =", en_cv.score(X_std, y))

en_best = ElasticNet(alpha=en_cv.alpha_, l1_ratio=en_cv.l1_ratio_, max_iter=10000)
en_best.fit(X_std, y)
en_coef_df = pd.DataFrame({'特征': feature_names, '系数': en_best.coef_})
print(en_coef_df.to_string(index=False))
""", "(10) 弹性网回归")

add_code_block(doc, r"""
# (11) 随机预留100个观测值作为测试集，岭回归
X_train, X_test, y_train, y_test = train_test_split(
    X_std, y, test_size=100, random_state=0
)

kf2 = KFold(n_splits=10, shuffle=True, random_state=1)
ridge_cv2 = RidgeCV(alphas=np.logspace(-3, 6, 100), cv=kf2, scoring='r2')
ridge_cv2.fit(X_train, y_train)

ridge_final = Ridge(alpha=ridge_cv2.alpha_)
ridge_final.fit(X_train, y_train)

print("最优 α =", ridge_cv2.alpha_)
print("训练集 R² =", round(ridge_final.score(X_train, y_train), 4))
print("测试集 R² =", round(ridge_final.score(X_test, y_test), 4))
""", "(11) 测试集评估（岭回归）")

# ---------- 保存 ----------
output_path = r'D:\桌面\student\惩罚回归代码.docx'
doc.save(output_path)
print(f"Word 文档已保存至: {output_path}")
